import os
import requests
import pandas as pd
import xml.etree.ElementTree as et
from flask import Blueprint, request, send_file, render_template
from io import BytesIO
from app.gift_fund_bibliography.auth_gift_fund_bibliography import login_required
from app.gift_fund_bibliography.citeproc_local.py2compat import *
from app.gift_fund_bibliography.html_helper import HTMLHelper


from app.gift_fund_bibliography.citeproc_local.source.bibtex import BibTeX
import unicodedata
from app.gift_fund_bibliography.citeproc_local import CitationStylesStyle, CitationStylesBibliography
from app.gift_fund_bibliography.citeproc_local import formatter
from app.gift_fund_bibliography.citeproc_local import Citation, CitationItem
from django.utils.encoding import python_2_unicode_compatible, smart_text, smart_bytes
import docx
import json
import re
from dotenv import load_dotenv
import zipfile
import io
import sys

load_dotenv()

gift_fund_blueprint = Blueprint("gift_fund_bibliography", __name__)


class GiftFundBibliography:
    def __init__(self, library, fiscal_year):
        self.library = library
        self.fiscal_year = fiscal_year
        self.word_docs = {}
        self.api_key = os.getenv("analytics_api_key")
        self.sru_url = "https://tufts.alma.exlibrisgroup.com/view/sru/01TUN_INST?version=1.2&operation=searchRetrieve&recordSchema=marcxml&query=alma.mms_id="
        self.mms_id_and_fund_df = None
        self.mms_id_list = []
        self.marc_df = None
        self.output_dir = "./Output"
        self.processing_dir = "./Processing"
        self.bib_buffers = {}
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.processing_dir, exist_ok=True)
        self.count_file = BytesIO()
        self.error_file = BytesIO()
        self.output_file = BytesIO()

    def process(self):
        report_content = self.retrieve_analytics_report()
        self.parse_analytics_report(report_content)
        self.retrieve_bib_records()
        self.merge_data()
        self.clean_data()
        self.create_bib_dataframes_and_buffers()
        self.generate_bibliography()
        return self.create_zip()
        
    
    # 
    def create_bib_dataframes_and_buffers(self):
        import re

        
        df = self.full_df.copy()
        df = df.applymap(lambda x: smart_bytes(x).decode("utf-8") if isinstance(x, str) else x)
        df = df.replace("nan", "", regex=True)
        df = df[df["Title"].str.isupper() == False]
        df = df.drop_duplicates(subset=["Title", "Author Name"], keep="first")
        df = df.rename(columns={"Fund Ledger Code": "Fund"})

        fund_list = df["Fund"].unique().tolist()

        for fund in fund_list:
            bib_buffer = io.StringIO()
            gfSegment = df[df["Fund"] == fund].reset_index(drop=True)
            count = len(gfSegment.index)

            for x in range(count):
                title = gfSegment.iloc[x]["Title"]
                title = re.sub(r"(^.+)\.$", r"\1", title).strip()
                title = re.sub(r"\/", "", title)
                
                if not title or gfSegment.iloc[x]["Title"].isupper():
                    continue

                test_mms_id = gfSegment.iloc[x]["MMS Id"]

                # Combine all authors/editors/translators into a single author field
                author_names = []
                for field in [
                    "Author Name", "Second Author Name", "Corporate Author Name", "Second Corporate Author Name"
                ]:
                    val = gfSegment.iloc[x].get(field, "")
                    val = val.replace(".", "")
                    if val and val != "Empty":
                        val = [auth.strip(" ,;") for auth in val.split(";") if auth.strip()]
                        author_names.extend(val)

                if author_names:
                    
                    author_field = f"\tauthor = {{{' and '.join(author_names)}}},\n"
                else:
                    author_field = ""


                format_note = ""
                format_field = gfSegment.iloc[x]["Format"]
                if re.search(r"[Ee]lectronic", format_field):
                    format_str = re.sub(r"^.*?([Ee]lectronic\s[A-Za-z- ]+)", r"\1", format_field)
                    format_str = re.sub(r"s\.$", "", format_str)
                    format_note = f"\tnote = {{<i>{format_str}</i>}},\n"

                # Clean year
                raw_year = gfSegment.iloc[x]["First Published Year"] or gfSegment.iloc[x]["Second Published Year"]
                if raw_year:
                    year = re.sub(r"[^0-9]", "", raw_year)
                    year = re.sub(r"^(\d{4}).*", r"\1", year)
                else:
                    year = ""

                publicationInfo = self.parsePublication(
                    gfSegment.iloc[x]["First Place of Publication"].replace(";;", ";"),
                    gfSegment.iloc[x]["First Publisher"],
                    year,
                    gfSegment.iloc[x]["Second Place of Publication"].replace(";;", ";"),
                    gfSegment.iloc[x]["Second Publisher"],
                    year
                    
                )


                bib_buffer.write(f"@BOOK")
                bib_buffer.write("{")
                bib_buffer.write(f"{test_mms_id},\n")
                
                if author_field:
                    bib_buffer.write(author_field)
                bib_buffer.write(f"\ttitle = {{{title}}},\n")
                bib_buffer.write(publicationInfo)
                bib_buffer.write(format_note)
                bib_buffer.write("}\n\n")

                # print(bib_buffer)

                

                
            self.bib_buffers[fund] = bib_buffer
  
        # print(bib_buffer)
    def retrieve_analytics_report(self):
        """Retrieve the analytics report using the API."""
        url = f"https://api-na.hosted.exlibrisgroup.com/almaws/v1/analytics/reports?apikey={self.api_key}"
        limit = "&limit=1000"
        format = "&format=xml"
        path = self._get_report_path()
        filter = self._get_report_filter()
        print(url + format + path + limit + filter)
        response = requests.get(url + format + path + limit + filter)
        
        report = requests.get(url + format + path + limit + filter)


        report_string = report.content
        # #
        print(str(report_string))
        #
        # file = open("test report string.txt", 'w+')
        #
        # file.write(str(report_string))
        # sys.exit()
        tree = et.ElementTree(et.fromstring(report.content))

        root = tree.getroot()

        isFinished = root[0][1].text
        resumptionToken = root[0][0].text
        isFinishedContinue = isFinished
        result_set = root[0][2][0]
        ##############################################################################################################################
        ##############################################################################################################################
        ##############################################################################################################################
        ########    Retrieve Analytics reports for initial input:
        ########    - list of MMS IDs from set list of gift funds purchased in the current fiscal year.
        ########      The list of funds will be be static and part of the Analytics report.
        ########      The date range will have to be passed as a SAW XML filter to the generic report, for Tisch.
        ########      The user will set the date range by entering fiscal year as "FY\d\d\d\d" in
        ########      the first prompt, and choose library by choosing either Tisch or Ginn in the second.
        ########      The library chosen will affect which report it goes to.
        ########    - The new version of this report will also contain fund, so that I don't need to retrieve both reports,
        ########      but for the intiial query it will just use MMS ID
        ######################################################################################################
        ######################################################################################################
        #######     since, output is limited to 1,000 records, iterate through whole report
        while isFinishedContinue == "false":
            full_path = url + "&token=" + resumptionToken
            reportContinue = requests.get(url + "&token=" + resumptionToken)
            # print ("\n\n\n\n" + str(reportContinue.content) + "\n\n\n\n")
            report_string += reportContinue.content

            treeContinue = et.ElementTree(et.fromstring(reportContinue.content))
            rootContinue = treeContinue.getroot()
            result_set_continue = rootContinue[0][1][0]
            result_set.append(result_set_continue)

            isFinishedContinue = rootContinue[0][0].text
        return result_set
    def _get_report_path(self):
        """Get the report path based on the library."""
        if self.library == "Tisch Library":
            return "&path=%2Fshared%2FTufts%20University%2FReports%2FCollections%2FGift%20Funds%2FTitles%20Purchased%20with%20Gift%20Funds%20-%20Tisch%20-%20Generic%20for%20Script"
        elif self.library == "Ginn Library":
            return "&path=%2Fshared%2FTufts%20University%2FReports%2FCollections%2FGift%20Funds%2FTitles%20Purchased%20with%20Gift%20Funds%20-%20Ginn%20-%20Generic%20for%20Script"

    def _get_report_filter(self):
        """Get the report filter based on the fiscal year."""
        return (
            "&filter=%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22and%22+xmlns%3Asaw%3D%22com.siebel.analytics.web%2Freport%2Fv1.1%22+%0D%0A++xmlns%3Asawx%3D%22com.siebel.analytics.web%2Fexpression%2Fv1.1%22+%0D%0A++xmlns%3Axsi%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema-instance%22+%0D%0A++xmlns%3Axsd%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema%22%3E%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22notEqual%22%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Bibliographic+Details%22.%22MMS+Id%22%3C%2Fsawx%3Aexpr%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Astring%22%3E-1%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%09%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22equal%22%3E%0D%0A+++++++++++++++%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Transaction+Date%22.%22Transaction+Date+Fiscal+Year%22%3C%2Fsawx%3Aexpr%3E%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Adecimal%22%3E"
            + str(self.fiscal_year)
            + "%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%3C%2Fsawx%3Aexpr%3E"
        )

    def parse_analytics_report(self, result_set):
        
        """Parse the analytics report to extract MMS IDs and fund information."""
        # tree = et.ElementTree(et.fromstring(report_content))
        # root = tree.getroot()
        # result_set = root[0][2][0]
        mms_id_and_fund_dict = []
        for element in result_set.iter():
            if re.match(r".*Row", element.tag):
                mms_id = ""
                fund_code = ""
                for subElement in element.iter():
                    if re.match(r".*Column1", subElement.tag):
                        mms_id = subElement.text
                        self.mms_id_list.append(mms_id)
                    elif re.match(r".*Column2", subElement.tag):
                        fund_code = subElement.text
                mms_id_and_fund_dict.append(
                    {"MMS Id": mms_id, "Fund Ledger Code": fund_code}
                )
        self.mms_id_and_fund_df = pd.DataFrame(mms_id_and_fund_dict)
        # self.mms_id_and_fund_df.to_csv(
        #     self.processing_dir + "/MMS ID and Fund.csv", index=False
        # )

    def retrieve_bib_records(self):
            sru_url = "https://tufts.alma.exlibrisgroup.com/view/sru/01TUN_INST?version=1.2&operation=searchRetrieve&recordSchema=marcxml&query=alma.mms_id="
            records = []
            namespaces = {"ns1": "http://www.loc.gov/MARC21/slim"}

            for mms_id in self.mms_id_list:
                record = {"MMS Id": mms_id}
                print(f"retrieve bib record for {mms_id}")
                try:
                    response = requests.get(sru_url + str(mms_id))
                    root = et.fromstring(response.content.decode("utf-8"))

                    def get_field(tag, code):
                        return [sub.text for sub in root.findall(f".//ns1:datafield[@tag='{tag}']/ns1:subfield[@code='{code}']", namespaces)]

                    def get_single(tag, code):
                        val = root.find(f".//ns1:datafield[@tag='{tag}']/ns1:subfield[@code='{code}']", namespaces)
                        return val.text if val is not None else ""

                    record["Author Name"] = ";".join(get_field("100", "a"))
                    record["Author Relator"] = ";".join(get_field("100", "e"))
                    record["Second Author Name"] = ";".join(get_field("700", "a"))
                    record["Second Author Relator"] = ";".join(get_field("700", "e"))
                    record["Corporate Author Name"] = ";".join(get_field("110", "a"))
                    record["Corporate Author Relator"] = ";".join(get_field("110", "e"))
                    record["Second Corporate Author Name"] = ";".join(get_field("710", "a"))
                    record["Second Corporate Author Relator"] = ";".join(get_field("710", "e"))
                    record["Format"] = next((sff for sff in get_field("655", "a") if "Electronic" in sff), "")
                    title_a = get_single("245", "a")
                    title_b = get_single("245", "b")
                    record["Title"] = (title_a or "") + (" " + title_b if title_b else "")
                    record["First Place of Publication"] = ";".join(get_field("260", "a"))
                    record["First Publisher"] = ";".join(get_field("260", "b"))
                    record["First Published Year"] = ";".join(get_field("260", "c"))
                    record["Second Place of Publication"] = ";".join(get_field("264", "a"))
                    record["Second Publisher"] = ";".join(get_field("264", "b"))
                    record["Second Published Year"] = ";".join(get_field("264", "c"))
                    records.append(record)
                except Exception as e:
                    self.error_file.write(f"Error retrieving/parsing MARC for {mms_id}: {e}\n")

            print(self.marc_df)
            self.marc_df = pd.DataFrame(records)

    def parse_bib_records(self, marc_records):
        records = []
        namespaces = {"ns1": "http://www.loc.gov/MARC21/slim"}
        for record_content in marc_records:
            tree = et.ElementTree(et.fromstring(record_content))
            root = tree.getroot()
            data = {"MMS Id": root.find(".//ns1:controlfield[@tag='001']", namespaces).text}

            fields = [
                ("Author Name", ".//ns1:datafield[@tag='100']/ns1:subfield[@code='a']"),
                ("Author Relator", ".//ns1:datafield[@tag='100']/ns1:subfield[@code='e']"),
                ("Second Author Name", ".//ns1:datafield[@tag='700']/ns1:subfield[@code='a']"),
                ("Second Author Relator", ".//ns1:datafield[@tag='700']/ns1:subfield[@code='e']"),
                ("Corporate Author Name", ".//ns1:datafield[@tag='110']/ns1:subfield[@code='a']"),
                ("Corporate Author Relator", ".//ns1:datafield[@tag='110']/ns1:subfield[@code='e']"),
                ("Second Corporate Author Name", ".//ns1:datafield[@tag='710']/ns1:subfield[@code='a']"),
                ("Second Corporate Author Relator", ".//ns1:datafield[@tag='710']/ns1:subfield[@code='e']"),
                ("Format", ".//ns1:datafield[@tag='655']/ns1:subfield[@code='a']"),
                ("Title", ".//ns1:datafield[@tag='245']/ns1:subfield[@code='a']"),
                ("First Place of Publication", ".//ns1:datafield[@tag='260']/ns1:subfield[@code='a']"),
                ("First Publisher", ".//ns1:datafield[@tag='260']/ns1:subfield[@code='b']"),
                ("First Published Year", ".//ns1:datafield[@tag='260']/ns1:subfield[@code='c']"),
                ("Second Place of Publication", ".//ns1:datafield[@tag='264']/ns1:subfield[@code='a']"),
                ("Second Publisher", ".//ns1:datafield[@tag='264']/ns1:subfield[@code='b']"),
                ("Second Published Year", ".//ns1:datafield[@tag='264']/ns1:subfield[@code='c']")
            ]

            for field_name, xpath in fields:
                elems = root.findall(xpath, namespaces)
                data[field_name] = ";".join([e.text for e in elems if e.text]) if elems else ""

            records.append(data)
        return pd.DataFrame(records)

    ######################################################################################################
    ######################################################################################################
    #######     create lists of creators (either author, editor, or translators)
    def parseCreatorList(cList, relator):
        count = len(cList)
        x = 0
        creatorLine = ""
        while x < count and x < 3:


            if "author" in relator:
                #print("In author loop\n")
                if x == 0 and count == 1:
                    creatorLine += "\tauthor = {" + cList[x] + "}"
                    break
                elif x == 0 and count > 1:
                    creatorLine += "\tauthor = {" + cList[x]
                elif 0 < x < 2 and count > 2:
                    creatorLine += " and " + cList[x]
                else:
                    creatorLine += " and " + cList[x] + "}"
            elif "editor" in relator:

                if x == 0 and count == 1:
                    creatorLine += "\teditor = {" + cList[x] + "}"
                    break
                elif x == 0 and count > 1:
                    creatorLine += "\teditor = {" + cList[x]
                elif 0 < x < 2 and count > 2:
                    creatorLine += " and " + cList[x]
                else:
                    creatorLine += " and " + cList[x] + "}"
            elif "translator" in relator:
                #print("In transalator loop\n")
                if x == 0 and count == 1:
                    creatorLine += "\ttranslator = {" + cList[x] + "}"
                    break
                elif x == 0 and count > 1:
                    creatorLine += "\ttranslator = {" + cList[x]
                elif 0 < x < 2 and count > 2:
                    creatorLine += " and " + cList[x]
                else:
                    creatorLine += " and " + cList[x] + "}"
            #if it's uncaught relator
            else:
                #print("In uncaught relator loop\n")
                if x == 0 and count == 1:
                    creatorLine += "\tauthor = {" + cList[x] + "}"
                    break
                elif x == 0 and count > 1:
                    creatorLine += "\tauthor = {" + cList[x]
                elif 0 < x < 2 and count > 2:
                    creatorLine += " and " + cList[x]
                else:
                    creatorLine += " and " + cList[x] + "}"

            x += 1

        return creatorLine

    ######################################################################################################
    ######################################################################################################
    #######     parse strings into lists of incoming creators
    def parseCreator(self, c, cR, type, mms_id):
        creatorFlag = False
        relatorFlag = False

        if c != "":
            cList = c.split(";")
            creatorFlag = True
        else:
            creatorFlag = False
        if cR != "":
            cRList = cR.split(";")
            relatorFlag = True
        else:
            relatorFlag = False


        authorList = []
        editorList = []
        translatorList = []
        y = 0
        nullVariable = ""
        if creatorFlag == True:
            for creator in cList:


                if type == "personal":
                    cList[y] = re.sub(r'([^,]+,\s[^,]+),', r'\1', cList[y])
                    cList[y] = re.sub(r'([^,.]+?)[,.]\W(.+),?', r'\2 \1', str(cList[y]))
                creator = cList[y]
                #if relatorFlag:
                if relatorFlag == True:
                    try:
                        relator = cRList[y]
                        if "author" in relator:
                            authorList.append(creator)
                        elif "editor" in relator:
                            editorList.append(creator)
                        elif "translator" in relator:
                            translatorList.append(creator)
                    except:
                        authorList.append(creator)
                else:
                    authorList.append(creator)
                #else:
                    #authorList.append(creator)
                y += 1




            returnCreator = ""

            authorLine = ""
            editorLine = ""
            translatorLine = ""


            if len(authorList) > 0:
                authorLine = self.parseCreatorList(authorList, "author")
            if len(editorList) > 0:
                editorLine = self.parseCreatorList(editorList, "editor")
            if len(translatorList) > 0:
                translatorLine = self.parseCreatorList(translatorList, "translator")

            if authorLine != "":
                returnCreator += authorLine + ",\n"
            if editorLine != "":
                returnCreator += editorLine  + ",\n"
            if translatorLine != "":
                returnCreator += translatorLine + ",\n"

            if type == "corporate":
                returnCreator = re.sub(r'([a-z]+\s+\=\s+)({.+?\})', r'\1{\2}', returnCreator)

            return returnCreator

    #####################################################################################################
    #####################################################################################################
    ######     parse strings into lists of incoming publication info
    def parsePublication(self, a1, a2, a3, b1, b2, b3):
        address = ""
        publisher = ""
        year = ""
        format = ""



        if a2 != "":
            a2 = a2.split(";")
            a2String = str(a2[0])
            a2String = re.sub(r',$\[\]', '', str(a2String))

            if a1 != "":
                a1 = a1.split(";")
                a1String = str(a1[0])
                a1String = re.sub(r'\s+\:.*$', '', str(a1String))
                address = "\taddress = {" + str(a1String) + "},\n"
                publisher = "\tpublisher = {" + str(a2String) + "},\n"
                if a3 != "":
                    a3 = a3.split(";")
                    a3String = str(a3[0])
                    a3String = re.sub(r'.*(\d{4}).*', r'\1', str(a3String))
                    a3String = re.sub(r'[\[\]]', '', str(a3String))
                    if re.match(r'^\d+$', a3String):
                        year = "\tyear = {" + str(a3String) + "},\n"

        elif b2 != "":
            b2 = b2.split(";")
            b2String = str(b2[0])
            b2String = re.sub(r'\s\:.*$\[\]', '', b2String)
            b2String = re.sub(r',$', '', b2String)
            if b1 != "":
                b1 = b1.split(";")
                b1String = str(b1[0])
                b1String = re.sub(r'\s+\:.*$', '', b1String)
                address = "\taddress = {" + str(b1String) + "},\n"
                publisher = "\tpublisher = {" + str(b2String) + "},\n"
                if b3 != "":
                    b3 = b3.split(";")
                    b3String = str(b3[0])
                    b3String = re.sub(r'.*(\d{4}).*', r"\1", b3String)
                    b3String = re.sub(r'[\[\]]', '', b3String)

                    if re.match(r'^\d+$', b3String):
                        year = "\tyear = {" + str(b3String) + "},\n"

        # publisher_for_format = re.sub(r'.+?\{(.+?)\}', r'\1', publisher)
        # address_for_format = re.sub(r'.+?\{(.+?)\}', r'\1', address)
        # year_for_format = re.sub(r'.+?\{(.+?)\}', r'\1', year)

        # if f != "":
        #     format = "\thowpublished = {" + publisher_for_format + ", " + address_for_format + ", " + str(year_for_format) + "[" + format + "].},"
        #     return_publisher = format
        # else:
        return_publisher = address + publisher + year


        return_publisher = return_publisher.replace(',,', ',')
        return_publisher = return_publisher.replace('[', '')
        return_publisher = return_publisher.replace(']', '')

        return return_publisher
   

    def sanitize_bibtex_field(self, text):
        """Sanitize input for BibTeX: escape diacritics and special characters with correct LaTeX formatting."""
        if not text:
            return ""

        decomposed = unicodedata.normalize("NFD", text)
        result = ""

        combining_map = {
            "\u0300": "\\`", "\u0301": "\\'", "\u0302": "\\^", "\u0303": "\\~",
            "\u0304": "\\=", "\u0306": "\\u", "\u0307": "\\.", "\u0308": '\\"',
            "\u030A": "\\r", "\u030B": "\\H", "\u030C": "\\v",
            "\u0327": "\\c", "\u0328": "\\k",
        }

        literal_map = {
            "đ": "\\dj{}", "Đ": "\\DJ{}", "ø": "\\o{}", "Ø": "\\O{}",
            "ł": "\\l{}", "Ł": "\\L{}", "ß": "{\\ss}", "æ": "\\ae{}", "Æ": "\\AE{}",
            "œ": "\\oe{}", "Œ": "\\OE{}", "–": "-", "—": "-", "“": "``", "”": "''",
            "‘": "`", "’": "'",
        }

        skip_next = False
        for i, char in enumerate(decomposed):
            if skip_next:
                skip_next = False
                continue

            if char in literal_map:
                result += literal_map[char]
                continue

            if i + 1 < len(decomposed) and unicodedata.combining(decomposed[i + 1]):
                base = char
                diacritic = decomposed[i + 1]
                latex = combining_map.get(diacritic)
                if latex:
                    result += f"{latex}{{{base}}}"
                    skip_next = True
                else:
                    result += base
            elif not unicodedata.combining(char):
                result += char

        specials = {
            "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
            "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\~{}",
            "^": r"\^{}",
        }

        result = re.sub(
            '|'.join(map(re.escape, specials)),
            lambda m: specials[m.group()],
            result
        )

        result = result.rstrip(".,;: ")

        return result



    # def parsePublication(self, first_place, first_publisher, first_year, second_place, second_publisher, second_year):
    #     pub = ""

        

    #     first_place = first_place.replace(" ;", ", ")
    #     first_place = first_place.replace(";", "")
    #     first_place = re.sub(r"[;:,]$", "", first_place)
    #     first_publisher = first_publisher.replace(" ;", ", ")
    #     first_publisher = first_publisher.replace(";", "")
    #     first_publisher = re.sub(r"[;:,]$", "", first_publisher)
        
    #     first_publisher
    #     print(first_place + " - " + first_publisher + " - " + first_year)
    #     if first_place:
    #         pub += f"\taddress = {{{first_place}}},\n"
    #     if first_publisher:
    #         pub += f"\tpublisher = {{{first_publisher}}},\n"
    #     if first_year:
    #         pub += f"\tyear = {{{first_year}}},\n"
    #     print("pub")
    #     print(pub)
    #     return pub

    def merge_data(self):
        self.full_df = pd.merge(self.mms_id_and_fund_df, self.marc_df, on="MMS Id", how="inner")
        print("merged data")
        print(self.full_df)
    def clean_data(self):
        self.full_df = self.full_df.applymap(lambda x: smart_bytes(x).decode("utf-8") if isinstance(x, str) else x)
        self.full_df.replace("nan", "", regex=True, inplace=True)
        self.full_df = self.full_df[self.full_df["Title"].str.isupper() == False]
        self.full_df.drop_duplicates(subset=["Title", "Author Name"], keep="first", inplace=True)

        print("clean data")

        print(self.full_df)
    def warn(self, citation_item):
        self.error_file.write(f"WARNING: Reference with key '{citation_item.key}' not found in the bibliography.\n")

    
    def generate_bibliography(self):
        for fund, buffer in self.bib_buffers.items():
            # try:
            buffer.seek(0)
            bib_text = buffer.read()  # <-- read safely now
            buffer = io.StringIO(bib_text)  # recreate to pass to BibTeX

            print("bib text")
            print(bib_text)
            print("type bib text")
            print(type(bib_text))
            bib_source = BibTeX(io.StringIO(bib_text))

            bib_style = CitationStylesStyle("chicago-annotated-bibliography", validate=True)
            bibliography = CitationStylesBibliography(bib_style, bib_source, formatter.html)
            

            

            doc = docx.Document()
            doc.add_heading("References", 0)
            html_helper = HTMLHelper()
            for item in bib_source:
                
                citation = Citation([CitationItem(item)])
                
                bibliography.register(citation)
                
                item_string = bibliography.cite(citation, self.warn)
                
                html_helper = HTMLHelper()
                bibliography.sort()
            for item in bibliography.bibliography():
                # print(item)
                # sys.exit()
                ######################################################################################################
                ######################################################################################################
                #######     take out extra characters in citation, that are artifacts of the citeproc citation
                #######     creation process with some of our bib records
                #######
                #######     Also make the editor label plural if there are multiple ("eds.")
                item = str(item)
                item = item.replace(", n.d..", "")
                item = item.replace(',,', ',')
                item = item.replace('..', '.')
                item = re.sub(r'([^<]+?and[^<]+?)(ed.)(\s+<i>)', r'\1eds.\3', item)
                item = item.replace(',.', '.')
                # if re.search(r'^.*?([Ee]lectronic.*?).*?$', item):
                    # citation_format = re.sub(r'^.*?\[([Ee]lectronic.*?\]).*?$', r'\1', item)
                    # item = item + "\n" + str(citation_format)
                    #
                    # print(item)
                    # # sys.exit()
                ######################################################################################################
                ######################################################################################################
                #######     turn HTML into document styling
                run_map = html_helper.html_to_run_map(str(item))



                par = doc.add_paragraph()

                html_helper.insert_runs_from_html_map(par, run_map)
                # run_map = html_helper.html_to_run_map(item_string)
                # par = doc.add_paragraph()
                # html_helper.insert_runs_from_html_map(par, run_map)
                # print("Citation text:", item_string)
                # print("Run map:", run_map)

                # html_helper.insert_runs_from_html_map(par, run_map)

                # print("Added paragraph text:", par.text)


            doc_buffer = io.BytesIO()
            # print(doc_buffer.getvalue())
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            print("Final document buffer length:", len(doc_buffer.getvalue()))

            self.word_docs[fund] = doc_buffer

            # except Exception as e:
            #     self.error_file.write(f"Error processing BibTeX for fund {fund}:\n{bib_text}\nException: {str(e)}\n".encode("utf-8"))

    def create_zip(self):
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr("Count File.txt", self.count_file.getvalue())
                zip_file.writestr("Errors.txt", self.error_file.getvalue())
                for fund, buffer in self.word_docs.items():
                    zip_file.writestr(f"{fund}.docx", buffer.getvalue())
            zip_buffer.seek(0)
            return zip_buffer

    # def _create_bib_file(self, fund, fund_marc_df):
    #     """Create a BibTeX file for a specific fund."""
    #     bib_content = ""
    #     for _, row in fund_marc_df.iterrows():
    #         bib_content += f"@BOOK{{{row['MMS Id']}},\n"
    #         bib_content += f"\ttitle = {{{row['Title']}}},\n"
    #         bib_content += f"\tauthor = {{{row['Author']}}},\n"
    #         bib_content += "}\n\n"
    #     self.output_file.write(bib_content.encode("utf-8"))
    #     self._generate_word_doc(fund, bib_content)

    # def _generate_word_doc(self, fund, bib_content):
    #     """Generate a Word document from a BibTeX file."""
    #     bib_source = BibTeX(BytesIO(bib_content.encode("utf-8")))
    #     bib_style = CitationStylesStyle(
    #         "chicago-annotated-bibliography", validate=False
    #     )
    #     bibliography = CitationStylesBibliography(bib_style, bib_source, formatter.html)
    #     doc = docx.Document()
    #     doc.add_heading("References", 0)
    #     for item in bib_source:
    #         citation = Citation([CitationItem(item)])
    #         bibliography.register(citation)
    #         item_string = bibliography.cite(citation, lambda x: None)
    #         par = doc.add_paragraph()
    #         par.add_run(item_string)
    #     docx_buffer = BytesIO()
    #     doc.save(docx_buffer)
    #     docx_buffer.seek(0)
    #     self.output_file.write(docx_buffer.getvalue())

    # def _create_zip(self):
    #     """Create a ZIP archive in memory."""
    #     zip_buffer = BytesIO()
    #     with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
    #         zip_file.writestr("Count File.txt", self.count_file.getvalue())
    #         zip_file.writestr("Errors.txt", self.error_file.getvalue())
    #         zip_file.writestr("Output File.docx", self.output_file.getvalue())
    #     zip_buffer.seek(0)
    #     return zip_buffer


# @gift_fund_blueprint.route("/", methods=["GET", "POST"])
# @login_required
# def index():
#     if request.method == "POST":
        
#         return send_file(
#             f"{bibliography.output_dir}/{library}.docx", as_attachment=True
#         )
#     return render_template("gift_fund_bibliography.html")
