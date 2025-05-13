import os
import requests
import pandas as pd
import xml.etree.ElementTree as et
from flask import Blueprint, request, send_file, render_template
from io import BytesIO
from app.bib_2_holdings_541.auth_541 import login_required
from citeproc_local.source.bibtex import BibTeX
from citeproc_local import CitationStylesStyle, CitationStylesBibliography, formatter
from citeproc_local import Citation, CitationItem
from django.utils.encoding import python_2_unicode_compatible, smart_text, smart_bytes
import docx
import json
import re
from dotenv import load_dotenv
import zipfile
import io

load_dotenv()

gift_fund_blueprint = Blueprint("gift_fund_bibliography", __name__)


class GiftFundBibliography:
    def __init__(self, library, fiscal_year):
        self.library = library
        self.fiscal_year = fiscal_year
        self.api_key = os.getenv("API_KEY")
        self.sru_url = "https://tufts.alma.exlibrisgroup.com/view/sru/01TUN_INST?version=1.2&operation=searchRetrieve&recordSchema=marcxml&query=alma.mms_id="
        self.mms_id_and_fund_df = None
        self.marc_df = None
        self.output_dir = "./Output"
        self.processing_dir = "./Processing"
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.processing_dir, exist_ok=True)
        self.count_file = BytesIO()
        self.error_file = BytesIO()
        self.output_file = BytesIO()

    def process(self):
        report_content = self.retrieve_analytics_report()
        self.parse_analytics_report(report_content)
        self.retrieve_bib_records(self)
        self.merge_data()
        self.clean_data()
        self.generate_bibliography()
        
    
    def create_bib_dataframes_and_buffers(self):
        import re

        self.bib_buffers = {}
        df = self.full_df.copy()
        df = df.applymap(lambda x: smart_bytes(x).decode("utf-8") if isinstance(x, str) else x)
        df = df.replace("nan", "", regex=True)
        df = df[df["Title"].str.isupper() == False]
        df = df.drop_duplicates(subset=["Title", "Author Name"], keep="first")
        df = df.rename(columns={"Fund Ledger Code": "Fund"})

        fund_list = df["Fund"].unique().tolist()

        for fund in fund_list:
            bib_buffer = io.StringIO()
            gfSegment = df[df["Fund"] == fund].reset_index(drop=True)
            count = len(gfSegment.index)

            for x in range(count):
                title = gfSegment.iloc[x]["Title"]
                title = re.sub(r"(^.+)\.$", r"\1", title).strip()
                if not title or gfSegment.iloc[x]["Title"].isupper():
                    continue

                test_mms_id = gfSegment.iloc[x]["MMS Id"]
                creator = ""
                if gfSegment.iloc[x]["Author Name"]:
                    creator += self.parseCreator(
                        gfSegment.iloc[x]["Author Name"],
                        gfSegment.iloc[x]["Author Relator"],
                        "personal",
                        test_mms_id
                    )
                if gfSegment.iloc[x]["Second Author Name"] != "Empty":
                    creator += self.parseCreator(
                        gfSegment.iloc[x]["Second Author Name"],
                        gfSegment.iloc[x]["Second Author Relator"],
                        "personal",
                        test_mms_id
                    )
                if gfSegment.iloc[x]["Corporate Author Name"] != "Empty":
                    creator += self.parseCreator(
                        gfSegment.iloc[x]["Corporate Author Name"],
                        gfSegment.iloc[x]["Corporate Author Relator"],
                        "corporate",
                        test_mms_id
                    )
                if gfSegment.iloc[x]["Second Corporate Author Name"] != "Empty":
                    creator += self.parseCreator(
                        gfSegment.iloc[x]["Second Corporate Author Name"],
                        gfSegment.iloc[x]["Second Corporate Author Relator"],
                        "corporate",
                        test_mms_id
                    )

                if re.search(r"(author.+?)(\r\n|\r|\n)\t+(author.+?)(\r\n|\r|\n)", creator):
                    creator = re.sub(r"(\tauthor.+?)(\r\n|\r|\n)(\t+author.+?)(\r\n|\r|\n)", r"\1\2", creator)

                format_field = gfSegment.iloc[x]["Format"]
                format_note = ""
                if re.search(r"[Ee]lectronic", format_field):
                    format_str = re.sub(r"^.*?([Ee]lectronic\s[A-Za-z- ]+)", r"\1", format_field)
                    format_str = re.sub(r"s\.$", "", format_str)
                    format_note = f"\tnote = {{<i>{format_str}</i>}},\n"

                publicationInfo = self.parsePublication(
                    gfSegment.iloc[x]["First Place of Publication"],
                    gfSegment.iloc[x]["First Publisher"],
                    gfSegment.iloc[x]["First Published Year"],
                    gfSegment.iloc[x]["Second Place of Publication"],
                    gfSegment.iloc[x]["Second Publisher"],
                    gfSegment.iloc[x]["Second Published Year"],
                )

                bib_buffer.write(f"@BOOK{{{gfSegment.iloc[x]['MMS Id']}},\n")
                bib_buffer.write(creator)
                if title.endswith(" /"):
                    title = title[:-2]
                bib_buffer.write(f"\ttitle = {{{title}}},\n")
                bib_buffer.write(publicationInfo)
                bib_buffer.write(format_note)
                bib_buffer.write("}\n\n")

            self.bib_buffers[fund] = io.BytesIO(bib_buffer.getvalue().encode("utf-8"))


    def process(self):
        report_content = self.retrieve_analytics_report()
        self.parse_analytics_report(report_content)
        self.parse_bib_records(self)
    def retrieve_analytics_report(self):
        """Retrieve the analytics report using the API."""
        url = f"https://api-na.hosted.exlibrisgroup.com/almaws/v1/analytics/reports?apikey={self.api_key}"
        limit = "&limit=1000"
        format = "&format=xml"
        path = self._get_report_path()
        filter = self._get_report_filter()
        response = requests.get(url + format + path + limit + filter)
        return response.content

    def _get_report_path(self):
        """Get the report path based on the library."""
        if self.library == "Tisch Library":
            return "&path=%2Fshared%2FTufts%20University%2FReports%2FCollections%2FGift%20Funds%2FTitles%20Purchased%20with%20Gift%20Funds%20-%20Tisch%20-%20Generic%20for%20Script"
        elif self.library == "Ginn Library":
            return "&path=%2Fshared%2FTufts%20University%2FReports%2FCollections%2FGift%20Funds%2FTitles%20Purchased%20with%20Gift%20Funds%20-%20Ginn%20-%20Generic%20for%20Script"

    def _get_report_filter(self):
        """Get the report filter based on the fiscal year."""
        return (
            "&filter=%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22and%22+xmlns%3Asaw%3D%22com.siebel.analytics.web%2Freport%2Fv1.1%22+%0D%0A++xmlns%3Asawx%3D%22com.siebel.analytics.web%2Fexpression%2Fv1.1%22+%0D%0A++xmlns%3Axsi%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema-instance%22+%0D%0A++xmlns%3Axsd%3D%22http%3A%2F%2Fwww.w3.org%2F2001%2FXMLSchema%22%3E%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22notEqual%22%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Bibliographic+Details%22.%22MMS+Id%22%3C%2Fsawx%3Aexpr%3E%0D%0A%09%09%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Astring%22%3E-1%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%09%0D%0A%09%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3Acomparison%22+op%3D%22equal%22%3E%0D%0A+++++++++++++++%3Csawx%3Aexpr+xsi%3Atype%3D%22sawx%3AsqlExpression%22%3E%22Transaction+Date%22.%22Transaction+Date+Fiscal+Year%22%3C%2Fsawx%3Aexpr%3E%3Csawx%3Aexpr+xsi%3Atype%3D%22xsd%3Adecimal%22%3E"
            + str(self.fiscal_year)
            + "%3C%2Fsawx%3Aexpr%3E%0D%0A%09%3C%2Fsawx%3Aexpr%3E%0D%0A%3C%2Fsawx%3Aexpr%3E"
        )

    def parse_analytics_report(self, report_content):
        """Parse the analytics report to extract MMS IDs and fund information."""
        tree = et.ElementTree(et.fromstring(report_content))
        root = tree.getroot()
        result_set = root[0][2][0]
        mms_id_and_fund_dict = []
        for element in result_set.iter():
            if re.match(r".*Row", element.tag):
                mms_id = ""
                fund_code = ""
                for subElement in element.iter():
                    if re.match(r".*Column1", subElement.tag):
                        mms_id = subElement.text
                    elif re.match(r".*Column2", subElement.tag):
                        fund_code = subElement.text
                mms_id_and_fund_dict.append(
                    {"MMS Id": mms_id, "Fund Ledger Code": fund_code}
                )
        self.mms_id_and_fund_df = pd.DataFrame(mms_id_and_fund_dict)
        # self.mms_id_and_fund_df.to_csv(
        #     self.processing_dir + "/MMS ID and Fund.csv", index=False
        # )

    def retrieve_bib_records(self):
            sru_url = "https://tufts.alma.exlibrisgroup.com/view/sru/01TUN_INST?version=1.2&operation=searchRetrieve&recordSchema=marcxml&query=alma.mms_id="
            records = []
            namespaces = {"ns1": "http://www.loc.gov/MARC21/slim"}

            for mms_id in self.mms_id_list:
                record = {"MMS Id": mms_id}
                try:
                    response = requests.get(sru_url + str(mms_id))
                    root = et.fromstring(response.content.decode("utf-8"))

                    def get_field(tag, code):
                        return [sub.text for sub in root.findall(f".//ns1:datafield[@tag='{tag}']/ns1:subfield[@code='{code}']", namespaces)]

                    def get_single(tag, code):
                        val = root.find(f".//ns1:datafield[@tag='{tag}']/ns1:subfield[@code='{code}']", namespaces)
                        return val.text if val is not None else ""

                    record["Author Name"] = ";".join(get_field("100", "a"))
                    record["Author Relator"] = ";".join(get_field("100", "e"))
                    record["Second Author Name"] = ";".join(get_field("700", "a"))
                    record["Second Author Relator"] = ";".join(get_field("700", "e"))
                    record["Corporate Author Name"] = ";".join(get_field("110", "a"))
                    record["Corporate Author Relator"] = ";".join(get_field("110", "e"))
                    record["Second Corporate Author Name"] = ";".join(get_field("710", "a"))
                    record["Second Corporate Author Relator"] = ";".join(get_field("710", "e"))
                    record["Format"] = next((sff for sff in get_field("655", "a") if "Electronic" in sff), "")
                    title_a = get_single("245", "a")
                    title_b = get_single("245", "b")
                    record["Title"] = (title_a or "") + (" " + title_b if title_b else "")
                    record["First Place of Publication"] = ";".join(get_field("260", "a"))
                    record["First Publisher"] = ";".join(get_field("260", "b"))
                    record["First Published Year"] = ";".join(get_field("260", "c"))
                    record["Second Place of Publication"] = ";".join(get_field("264", "a"))
                    record["Second Publisher"] = ";".join(get_field("264", "b"))
                    record["Second Published Year"] = ";".join(get_field("264", "c"))
                    records.append(record)
                except Exception as e:
                    self.error_file.write(f"Error retrieving/parsing MARC for {mms_id}: {e}\n")

            self.marc_df = pd.DataFrame(records)

    def parse_bib_records(self, marc_records):
        records = []
        namespaces = {"ns1": "http://www.loc.gov/MARC21/slim"}
        for record_content in marc_records:
            tree = et.ElementTree(et.fromstring(record_content))
            root = tree.getroot()
            data = {"MMS Id": root.find(".//ns1:controlfield[@tag='001']", namespaces).text}

            fields = [
                ("Author Name", ".//ns1:datafield[@tag='100']/ns1:subfield[@code='a']"),
                ("Author Relator", ".//ns1:datafield[@tag='100']/ns1:subfield[@code='e']"),
                ("Second Author Name", ".//ns1:datafield[@tag='700']/ns1:subfield[@code='a']"),
                ("Second Author Relator", ".//ns1:datafield[@tag='700']/ns1:subfield[@code='e']"),
                ("Corporate Author Name", ".//ns1:datafield[@tag='110']/ns1:subfield[@code='a']"),
                ("Corporate Author Relator", ".//ns1:datafield[@tag='110']/ns1:subfield[@code='e']"),
                ("Second Corporate Author Name", ".//ns1:datafield[@tag='710']/ns1:subfield[@code='a']"),
                ("Second Corporate Author Relator", ".//ns1:datafield[@tag='710']/ns1:subfield[@code='e']"),
                ("Format", ".//ns1:datafield[@tag='655']/ns1:subfield[@code='a']"),
                ("Title", ".//ns1:datafield[@tag='245']/ns1:subfield[@code='a']"),
                ("First Place of Publication", ".//ns1:datafield[@tag='260']/ns1:subfield[@code='a']"),
                ("First Publisher", ".//ns1:datafield[@tag='260']/ns1:subfield[@code='b']"),
                ("First Published Year", ".//ns1:datafield[@tag='260']/ns1:subfield[@code='c']"),
                ("Second Place of Publication", ".//ns1:datafield[@tag='264']/ns1:subfield[@code='a']"),
                ("Second Publisher", ".//ns1:datafield[@tag='264']/ns1:subfield[@code='b']"),
                ("Second Published Year", ".//ns1:datafield[@tag='264']/ns1:subfield[@code='c']")
            ]

            for field_name, xpath in fields:
                elems = root.findall(xpath, namespaces)
                data[field_name] = ";".join([e.text for e in elems if e.text]) if elems else ""

            records.append(data)
        return pd.DataFrame(records)

    def parseCreator(self, name, relator, creator_type, mms_id):
        if not name:
            return ""
        entry = ""
        if creator_type == "personal":
            entry = f"\tauthor = {{{name}}},\n"
        elif creator_type == "corporate":
            entry = f"\tauthor = {{{name}}},\n"
        if relator and re.search(r'edit', relator, re.IGNORECASE):
            entry = f"\teditor = {{{name}}},\n"
        if relator and re.search(r'translat', relator, re.IGNORECASE):
            entry = f"\ttranslator = {{{name}}},\n"
        return entry

    def parsePublication(self, first_place, first_publisher, first_year, second_place, second_publisher, second_year):
        pub = ""
        if first_place:
            pub += f"\taddress = {{{first_place}}},\n"
        if first_publisher:
            pub += f"\tpublisher = {{{first_publisher}}},\n"
        if first_year:
            pub += f"\tyear = {{{first_year}}},\n"
        return pub

    def merge_data(self):
        self.full_df = pd.merge(self.analytics_df, self.marc_df, on="MMS Id", how="inner")

    def clean_data(self):
        self.full_df = self.full_df.applymap(lambda x: smart_bytes(x).decode("utf-8") if isinstance(x, str) else x)
        self.full_df.replace("nan", "", regex=True, inplace=True)
        self.full_df = self.full_df[self.full_df["Title"].str.isupper() == False]
        self.full_df.drop_duplicates(subset=["Title", "Author Name"], keep="first", inplace=True)

    def warn(self, citation_item):
        self.error_file.write(f"WARNING: Reference with key '{citation_item.key}' not found in the bibliography.\n")

    def generate_bibliography(self):
        grouped = self.full_df.groupby("Fund Ledger Code")
        for fund, group in grouped:
            bib_content = ""
            for _, row in group.iterrows():
                try:
                    title = row["Title"].strip().rstrip('.')
                    if not title or title.isupper():
                        continue
                    creators = ""
                    creators += self.parseCreator(row.get("Author Name", ""), row.get("Author Relator", ""), "personal", row["MMS Id"])
                    creators += self.parseCreator(row.get("Second Author Name", ""), row.get("Second Author Relator", ""), "personal", row["MMS Id"])
                    creators += self.parseCreator(row.get("Corporate Author Name", ""), row.get("Corporate Author Relator", ""), "corporate", row["MMS Id"])
                    creators += self.parseCreator(row.get("Second Corporate Author Name", ""), row.get("Second Corporate Author Relator", ""), "corporate", row["MMS Id"])
                    publication = self.parsePublication(
                        row.get("First Place of Publication", ""),
                        row.get("First Publisher", ""),
                        row.get("First Published Year", ""),
                        row.get("Second Place of Publication", ""),
                        row.get("Second Publisher", ""),
                        row.get("Second Published Year", "")
                    )
                    note = ""
                    if re.search(r"[Ee]lectronic", row.get("Format", "")):
                        format_str = re.sub(r"^.*?([Ee]lectronic\s[^; ]+?)", r". \1", row["Format"])
                        format_str = re.sub(r"s\.$", "", format_str)
                        note = f"\tnote = {{<i>{format_str}</i>}},\n"
                    bib_content += f"@BOOK{{{row['MMS Id']}},\n"
                    bib_content += creators
                    bib_content += f"\ttitle = {{{title}}},\n"
                    bib_content += publication
                    bib_content += note
                    bib_content += "}\n\n"
                except Exception as e:
                    
                    self.error_file.write(f"Error writing bibliography to doc for {row['MMS Id']}: {e}\n")
            if not bib_content:
                continue

            bib_source = BibTeX(io.BytesIO(bib_content.encode("utf-8")))
            bib_style = CitationStylesStyle("chicago-annotated-bibliography", validate=False)
            bibliography = CitationStylesBibliography(bib_style, bib_source, formatter.html)

            doc = docx.Document()
            doc.add_heading("References", 0)
            for item in bib_source:
                citation = Citation([CitationItem(item)])
                bibliography.register(citation)
                item_string = bibliography.cite(citation, self.warn)
                run_map = formatter.html_to_run_map(item_string)
                par = doc.add_paragraph()
                formatter.insert_runs_from_html_map(par, run_map)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            self.word_docs[fund] = buffer

    def create_zip(self):
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.writestr("Count File.txt", self.count_file.getvalue())
            zip_file.writestr("Errors.txt", self.error_file.getvalue())
            for fund, buffer in self.word_docs.items():
                zip_file.writestr(f"{fund}.docx", buffer.getvalue())
        zip_buffer.seek(0)
        return zip_buffer

    def _create_bib_file(self, fund, fund_marc_df):
        """Create a BibTeX file for a specific fund."""
        bib_content = ""
        for _, row in fund_marc_df.iterrows():
            bib_content += f"@BOOK{{{row['MMS Id']}},\n"
            bib_content += f"\ttitle = {{{row['Title']}}},\n"
            bib_content += f"\tauthor = {{{row['Author']}}},\n"
            bib_content += "}\n\n"
        self.output_file.write(bib_content.encode("utf-8"))
        self._generate_word_doc(fund, bib_content)

    def _generate_word_doc(self, fund, bib_content):
        """Generate a Word document from a BibTeX file."""
        bib_source = BibTeX(BytesIO(bib_content.encode("utf-8")))
        bib_style = CitationStylesStyle(
            "chicago-annotated-bibliography", validate=False
        )
        bibliography = CitationStylesBibliography(bib_style, bib_source, formatter.html)
        doc = docx.Document()
        doc.add_heading("References", 0)
        for item in bib_source:
            citation = Citation([CitationItem(item)])
            bibliography.register(citation)
            item_string = bibliography.cite(citation, lambda x: None)
            par = doc.add_paragraph()
            par.add_run(item_string)
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        self.output_file.write(docx_buffer.getvalue())

    # def _create_zip(self):
    #     """Create a ZIP archive in memory."""
    #     zip_buffer = BytesIO()
    #     with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
    #         zip_file.writestr("Count File.txt", self.count_file.getvalue())
    #         zip_file.writestr("Errors.txt", self.error_file.getvalue())
    #         zip_file.writestr("Output File.docx", self.output_file.getvalue())
    #     zip_buffer.seek(0)
    #     return zip_buffer


@gift_fund_blueprint.route("/", methods=["GET", "POST"])
@login_required
def index():
    if request.method == "POST":
        library = request.form.get("library")
        fiscal_year = request.form.get("fiscal_year")
        if not library or not fiscal_year:
            return "Library and Fiscal Year are required", 400
        bibliography = GiftFundBibliography(library, fiscal_year)
        zip_file = bibliography.process()
        report_content = bibliography.retrieve_analytics_report()
        bibliography.parse_analytics_report(report_content)
        bibliography.retrieve_bib_records()
        zip_file = bibliography.generate_bibliography()
        return send_file(
            zip_file,
            mimetype="application/zip",
            as_attachment=True,
            download_name="gift_fund_bibliography.zip",
        )
        return send_file(
            f"{bibliography.output_dir}/{library}.docx", as_attachment=True
        )
    return render_template("gift_fund_bibliography.html")
